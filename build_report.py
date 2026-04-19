"""
Build the PA1 research report as a Word .docx file.
Run after pltr_pa1_analysis.py so the figures exist in ./figures/.
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = "PLTR_PA1_research_report.docx"

doc = Document()

# Default font sizing
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)


def h1(text):
    p = doc.add_heading(text, level=1)
    return p


def h2(text):
    p = doc.add_heading(text, level=2)
    return p


def para(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p


def bullet(text):
    p = doc.add_paragraph(text, style="List Bullet")
    return p


def fig(path, caption, width_in=5.5):
    doc.add_picture(path, width=Inches(width_in))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True
    cap.paragraph_format.space_after = Pt(12)


# --------------------------------------------------------------------------- #
# Title block
# --------------------------------------------------------------------------- #
title = doc.add_heading("MSDS 451 Programming Assignment 1", level=0)
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run("Predicting Next-Day Direction of Palantir (PLTR) Returns "
                  "with Engineered Lag Features and XGBoost")
run.bold = True
run.font.size = Pt(13)

byline = doc.add_paragraph()
byline.alignment = WD_ALIGN_PARAGRAPH.CENTER
byline.add_run("Nadia Shiroglazova   |   Northwestern University, MSDS 451   "
               "|   April 2026").italic = True

# --------------------------------------------------------------------------- #
# 1. Problem Description
# --------------------------------------------------------------------------- #
h1("1. Problem Description")
para(
    "This project replicates the methodology from the WTI jump-start example "
    "on a different asset of personal interest: Palantir Technologies (ticker "
    "PLTR). Palantir is a U.S. data-analytics and AI-platform company that has "
    "become one of the most-watched names in the post-2023 generative-AI cycle. "
    "Its share price has been volatile and frequently driven by news flow, "
    "which makes it an interesting candidate for short-horizon directional "
    "forecasting: if any structure exists in lagged price/volume features, an "
    "asset like PLTR is plausibly where it would show up."
)
para(
    "The forecasting task is binary classification: given trading-day t, predict "
    "whether the log return on day t+1 will be strictly positive (Target = 1) "
    "or zero/negative (Target = 0). The motivation is twofold:"
)
bullet("(a) Pedagogical — practice the same feature-engineering, "
       "subset-selection, time-series cross-validation, and gradient-boosted-"
       "tree workflow demonstrated for WTI.")
bullet("(b) Personal — gauge whether a small set of price/volume features "
       "carries any directional signal for a stock the author is considering "
       "buying.")
para(
    "The technical-report standard from Miller (2025) sets the methodology: "
    "engineer 15 lag/EMA-based features, run all-possible-subsets logistic "
    "regression with AIC for feature selection, train an XGBoost classifier "
    "under TimeSeriesSplit cross-validation, tune via randomized search, and "
    "evaluate with ROC AUC, a confusion matrix, and a classification report."
)

# --------------------------------------------------------------------------- #
# 2. Data Preparation and Pipeline
# --------------------------------------------------------------------------- #
h1("2. Data Preparation and Pipeline")
para(
    "Daily OHLCV data for PLTR were retrieved from Yahoo! Finance via the "
    "yfinance package. The download window starts at the IPO date "
    "(2020-09-30) and runs through 2026-04-19, yielding 1,393 raw trading "
    "days. After computing the lag and EMA warm-up rows and dropping nulls, "
    "1,390 rows remain for modeling. Compared with WTI's 5,000+ training "
    "rows, the PLTR sample is roughly four times smaller — a meaningful "
    "limitation that informs how much we can trust any out-of-sample "
    "result."
)
para(
    "All feature engineering and DataFrame operations use Polars. The "
    "engineered feature set mirrors the jump-start specification exactly so "
    "that results are directly comparable:"
)
bullet("Lagged closing prices: CloseLag1, CloseLag2, CloseLag3.")
bullet("Daily intraday range and its lags: HML = High − Low, with HMLLag1, "
       "HMLLag2, HMLLag3.")
bullet("Daily open-to-close move and its lags: OMC = Open − Close, with "
       "OMCLag1, OMCLag2, OMCLag3.")
bullet("Lagged volume: VolumeLag1, VolumeLag2, VolumeLag3.")
bullet("Exponential moving averages of CloseLag1 (i.e. computed from the "
       "previous day's close to avoid leakage): CloseEMA2, CloseEMA4, "
       "CloseEMA8 — half-lives of 1, 2, and 4 days.")
para(
    "The continuous response is the daily log return, "
    "LogReturn = ln(Close_t / Close_{t-1}). The binary classification "
    "target is Target = 1 when LogReturn > 0 and Target = 0 otherwise."
)
para(
    "Class balance was checked before modeling: 704 up days (50.6%) and "
    "686 even/down days (49.4%). The classes are essentially balanced, so "
    "no resampling was applied. All 15 features were standardized with "
    "scikit-learn's StandardScaler before being passed to either the "
    "logistic regression used in subset selection or the XGBoost classifier "
    "used downstream."
)

# --------------------------------------------------------------------------- #
# 3. Research Design
# --------------------------------------------------------------------------- #
h1("3. Research Design")
h2("3.1 Feature selection — all possible subsets with AIC")
para(
    "Following Miller (2025), feature selection was performed by exhaustive "
    "enumeration: every non-empty subset of the 15 candidate features "
    "(2^15 − 1 = 32,767 subsets) was fit to a logistic regression on the "
    "binary target, and the Akaike Information Criterion was computed as "
    "AIC = 2k − 2 ln L̂, where k is the number of model parameters "
    "(features + intercept) and L̂ is the maximized likelihood. Lower AIC "
    "is preferred — it rewards goodness of fit and penalizes complexity."
)
para(
    "The ten lowest-AIC subsets were collected, and the features that "
    "appeared most often across those ten subsets were retained for "
    "downstream modeling. This recovers a small, robust feature subset "
    "without committing to any single model's idiosyncrasies."
)
h2("3.2 Cross-validation design")
para(
    "Because trading-day observations are temporally dependent, ordinary "
    "k-fold cross-validation would leak future information into training "
    "folds. Scikit-learn's TimeSeriesSplit was used instead, with "
    "n_splits = 5 and a 10-day gap between each training fold and the "
    "following test fold. The gap absorbs the longest EMA half-life "
    "(CloseEMA8, half-life of 4 days) plus a buffer, so test-set features "
    "cannot quietly carry information from training-set days."
)
h2("3.3 Modeling algorithm and hyperparameter search")
para(
    "XGBoost (XGBClassifier with binary:logistic objective) was the chosen "
    "estimator. A baseline run with n_estimators = 1,000 and otherwise "
    "default hyperparameters was evaluated under TimeSeriesSplit to "
    "establish a starting point. Then a RandomizedSearchCV with 100 "
    "iterations sampled from the same five-hyperparameter space used in the "
    "jump-start (max_depth, min_child_weight, subsample, learning_rate, "
    "n_estimators) under the same TimeSeriesSplit. The model returned by "
    "RandomizedSearchCV was refit on the full sample and used for the final "
    "evaluation."
)

# --------------------------------------------------------------------------- #
# 4. Results
# --------------------------------------------------------------------------- #
h1("4. Results")
h2("4.1 Selected feature subset")
para(
    "The ten lowest-AIC subsets were dominated by intraday-range and "
    "intraday-direction features. Frequency counts across the top ten:"
)
bullet("HMLLag2 — appeared in 10 of 10 subsets")
bullet("HMLLag1 — appeared in 8 of 10")
bullet("OMCLag3 — appeared in 4 of 10")
bullet("VolumeLag2 — appeared in 4 of 10")
bullet("OMCLag1 — appeared in 3 of 10")
para(
    "These five features were carried forward. This is a striking departure "
    "from the WTI result reported in the jump-start, where the selected "
    "subset was {CloseLag3, HMLLag1, OMCLag2, OMCLag3, CloseEMA8} — i.e., "
    "anchored by lagged closes and the long EMA. For PLTR, none of "
    "CloseLag1–3 nor any of the three EMAs survived selection. The "
    "correlation heat map (Figure 1) suggests why: HMLLag1, HMLLag2, and "
    "CloseLag1 are correlated at r ≈ 0.80, so the intraday-range features "
    "absorb most of the autoregressive information that lagged closes would "
    "have contributed, and AIC's complexity penalty then prefers the "
    "smaller HML-anchored set. Substantively, this fits PLTR's character "
    "as a high-volatility, news-driven name where the daily true range "
    "carries more information than the price level itself."
)
fig("figures/pltr_correlation_heatmap.png",
    "Figure 1. Correlation heat map for the five selected features, plus "
    "LogReturn and CloseLag1.")

para(
    "As in the WTI case, every selected feature has a near-zero correlation "
    "with LogReturn (|r| ≤ 0.02). This forecasts up front that the "
    "downstream classifier will struggle on out-of-sample test folds, even "
    "if it can fit the training data nearly perfectly."
)

h2("4.2 Cross-validated baseline")
para(
    "The baseline XGBoost classifier (n_estimators = 1,000, defaults "
    "elsewhere) achieved a five-fold time-series CV accuracy of "
    "0.506 ± 0.036 — essentially indistinguishable from the 50.6% rate of "
    "the majority class. This is the honest signal in the data: with only "
    "five lagged features and ~1,400 days of history, the model has "
    "approximately no edge over a coin flip on truly held-out days."
)

h2("4.3 Hyperparameter tuning")
para(
    "RandomizedSearchCV (100 iterations, TimeSeriesSplit gap = 10, "
    "n_splits = 5) returned the following best settings:"
)
bullet("max_depth: 6")
bullet("min_child_weight: 1")
bullet("subsample: 0.893")
bullet("learning_rate: 0.014")
bullet("n_estimators: 846")
para(
    "Best cross-validated accuracy at these settings was 0.521 — a small "
    "improvement over the baseline, but still well within the noise band."
)

h2("4.4 Final-model evaluation on the full sample")
para(
    "Following the jump-start convention, the final model was refit on the "
    "full sample with the tuned hyperparameters and evaluated against the "
    "training labels. In-sample ROC AUC was 0.997 (Figure 2). The "
    "confusion matrix (Figure 3) shows 669 true negatives, 683 true "
    "positives, 17 false positives, and 21 false negatives — overall "
    "in-sample accuracy of 97.3%."
)
fig("figures/pltr_roc_curve.png",
    "Figure 2. In-sample ROC curve for the tuned XGBoost model. The very "
    "high AUC (0.997) reflects training-set fit, not held-out performance.")
fig("figures/pltr_confusion_matrix.png",
    "Figure 3. In-sample confusion matrix for the tuned XGBoost model.")

para("Per-class metrics (in-sample classification report):")
bullet("Negative return: precision 0.97, recall 0.98, F1 0.97 (n = 686)")
bullet("Positive return: precision 0.98, recall 0.97, F1 0.97 (n = 704)")

para(
    "The contrast between the cross-validated accuracy (~0.52) and the "
    "in-sample AUC (0.997) is the headline finding of this study. With a "
    "max_depth of 6 and 846 trees, XGBoost can memorize 1,390 standardized "
    "feature vectors essentially perfectly even when the features carry "
    "almost no genuine signal about the target. This is a textbook example "
    "of why time-series cross-validation matters more than training-set "
    "metrics when assessing a financial forecasting model."
)

fig("figures/pltr_feature_importance.png",
    "Figure 4. XGBoost gain-based feature importances on the final model.")

# --------------------------------------------------------------------------- #
# 5. Conclusions and Next Steps
# --------------------------------------------------------------------------- #
h1("5. Conclusions and Next Steps")
para(
    "Three takeaways from applying the jump-start methodology to PLTR:"
)
bullet(
    "Feature selection is asset-specific. PLTR's AIC-best subset is built "
    "around intraday range (HMLLag1, HMLLag2) rather than lagged closes "
    "and long EMAs. Practitioners should not assume that a feature subset "
    "tuned on one asset transfers to another."
)
bullet(
    "Honest out-of-sample performance is roughly chance. Time-series "
    "cross-validated accuracy hovers near 52%, only marginally above the "
    "majority-class baseline. The training-set AUC of 0.997 is misleading "
    "and demonstrates the importance of disciplined evaluation."
)
bullet(
    "PLTR's short trading history limits what can be learned. With ~1,400 "
    "trading days versus WTI's 5,000+, both feature selection and "
    "hyperparameter tuning sit on a smaller statistical foundation."
)
para("Plausible next steps include:")
bullet("Adding non-price features that may actually carry directional "
       "information for PLTR specifically — sector ETF returns (e.g. XLK), "
       "volatility-index levels (VIX), Treasury-yield changes, or sentiment "
       "scores from news headlines and earnings-call transcripts.")
bullet("Redefining the target as a three-level outcome (clear up / "
       "stable / clear down) tied to a magnitude threshold, so the model "
       "can decline to trade on small moves rather than being forced into a "
       "binary call.")
bullet("Holding out a true forward test set (e.g. the most recent 6–12 "
       "months) and reporting metrics there, rather than on the training "
       "data, before treating any model output as actionable.")
bullet("Repeating the workflow on related tickers (e.g. C3.AI, "
       "Snowflake, Datadog) to test whether the HML-dominant feature "
       "pattern generalizes across AI/data-platform names.")

# --------------------------------------------------------------------------- #
# 6. References
# --------------------------------------------------------------------------- #
h1("6. References")
bullet("Miller, T. W. (2025). 451 Feature Engineering: Programming "
       "Assignment 1 — Technical Report. Northwestern University, MSDS 451.")
bullet("Akaike, H. (1974). A new look at the statistical model "
       "identification. IEEE Transactions on Automatic Control, 19(6), "
       "716–723.")
bullet("Chen, T., & Guestrin, C. (2016). XGBoost: A scalable tree boosting "
       "system. Proceedings of KDD '16, 785–794.")
bullet("Pedregosa, F. et al. (2011). Scikit-learn: Machine learning in "
       "Python. JMLR, 12, 2825–2830.")
bullet("Géron, A. (2022). Hands-On Machine Learning with Scikit-Learn, "
       "Keras, and TensorFlow (3rd ed.). O'Reilly Media.")
bullet("yfinance Python package — https://github.com/ranaroussi/yfinance.")
bullet("Polars Python documentation — https://docs.pola.rs/.")

doc.save(OUT)
print(f"Wrote {OUT}")
